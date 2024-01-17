from docx import Document

class Word:
    def __init__(self, name):
        self.name = name
        self.fayl = Document()

    def read_text(self):
        text = ""
        for paragraph in self.fayl.paragraphs:
            text += paragraph.text + "\n"
        return text

    def add_text(self, paragraph):
        self.fayl.add_paragraph(paragraph)

    def add_heading(self, heading):
        self.fayl.add_heading(heading)

    def first_paragraph(self):
        return self.fayl.paragraphs[0].text

    def massiv(self):
        a = []
        for i in self.fayl.paragraphs:
            a.append(i.text)
        return a

    def __repr__(self):
        return self.name

    def __call__(self, n):
        return self.fayl.paragraphs[n].text

    def save(self, n):
        self.fayl.save(n)